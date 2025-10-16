import pdfplumber
import pandas as pd
import re
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches

# -------------------------------------------------
# CONFIGURATION
# -------------------------------------------------
books = [
    r"Livres_HarryPotter\Harry1.pdf", r"Livres_HarryPotter\Harry2.pdf", r"Livres_HarryPotter\Harry3.pdf",
    r"Livres_HarryPotter\Harry4.pdf", r"Livres_HarryPotter\Harry5.pdf", r"Livres_HarryPotter\Harry6.pdf", r"Livres_HarryPotter\Harry7.pdf"
]

characters = ["Harry", "Hermione", "Ron", "Rogue", "Dumbledore"]
illegal_words = ["vol", "mensonge", "crime", "triche", "magie noire", "violence", "harcèlement"]

# -------------------------------------------------
# INITIALISATION DU DATAFRAME
# -------------------------------------------------
stats = pd.DataFrame(columns=[
    "Livre", "Pages", "Harry_cicatrice", "Hermione_mais",
    "Harry_dialogue", "Hermione_dialogue", "Ron_dialogue",
    "Rogue_dark", "Dumbledore_intervention", "Actes_illegaux"
])

# -------------------------------------------------
# FONCTION DE COMPTE PAR REGEX
# -------------------------------------------------
def count_regex(pattern, text):
    return len(re.findall(pattern, text, re.IGNORECASE))

# -------------------------------------------------
# TRAITEMENT DES LIVRES
# -------------------------------------------------
for book in books:
    harry_cicatrice = 0
    hermione_mais = 0
    harry_dialogue = 0
    hermione_dialogue = 0
    ron_dialogue = 0
    rogue_dark = 0
    dumbledore_intervention = 0
    actes_illegaux = 0
    total_pages = 0

    try:
        with pdfplumber.open(book) as pdf:
            total_pages = len(pdf.pages)
            for page_number, page in enumerate(pdf.pages, start=1):
                try:
                    text = page.extract_text()
                    if text:
                        harry_cicatrice += count_regex(r"\bHarry\b.*\bcicatrice\b", text)
                        hermione_mais += count_regex(r"Hermione.*\bmais\b", text)
                        harry_dialogue += count_regex(r"Harry dit", text)
                        hermione_dialogue += count_regex(r"Hermione dit", text)
                        ron_dialogue += count_regex(r"Ron dit", text)
                        rogue_dark += count_regex(r"Rogue.*\b(mystérieux|dark)\b", text)
                        dumbledore_intervention += count_regex(r"Dumbledore.*(ordonne|demande|intervient|dirige|impose|tourne et retourne la chemise)", text)
                        actes_illegaux += sum([count_regex(word, text) for word in illegal_words])
                except Exception as e:
                    print(f"⚠️ Problème sur {book}, page {page_number} : {e}")
                    continue

        stats = pd.concat([stats, pd.DataFrame([{
            "Livre": book,
            "Pages": total_pages,
            "Harry_cicatrice": harry_cicatrice,
            "Hermione_mais": hermione_mais,
            "Harry_dialogue": harry_dialogue,
            "Hermione_dialogue": hermione_dialogue,
            "Ron_dialogue": ron_dialogue,
            "Rogue_dark": rogue_dark,
            "Dumbledore_intervention": dumbledore_intervention,
            "Actes_illegaux": actes_illegaux
        }])], ignore_index=True)

    except Exception as e:
        print(f"❌ Impossible d’ouvrir le fichier {book} : {e}")
        continue

# -------------------------------------------------
# NORMALISATION PAR 100 PAGES
# -------------------------------------------------
for col in ["Harry_cicatrice", "Hermione_mais", "Harry_dialogue", "Hermione_dialogue",
            "Ron_dialogue", "Rogue_dark", "Dumbledore_intervention", "Actes_illegaux"]:
    stats[f"{col}_par100pages"] = stats[col] / stats["Pages"] * 100

# -------------------------------------------------
# SAUVEGARDE CSV
# -------------------------------------------------
stats.to_csv("HarryPotter_stats.csv", index=False)
print("✅ Statistiques exportées dans HarryPotter_stats.csv")

# -------------------------------------------------
# VISUALISATIONS AMÉLIORÉES
# -------------------------------------------------
sns.set(style="whitegrid")
stats_plot = stats.set_index("Livre")

# 1️⃣ Dialogue des personnages
plt.figure(figsize=(14,7))
stats_plot[["Harry_dialogue","Hermione_dialogue","Ron_dialogue"]].plot(
    kind="bar", stacked=False, figsize=(14,7), colormap='Set2'
)
plt.title("Nombre de dialogues par personnage par livre")
plt.ylabel("Occurrences")
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig("dialogues_par_livre.png")
plt.show()

# 2️⃣ Actions spéciales et actes répréhensibles
plt.figure(figsize=(14,7))
stats_plot[["Harry_cicatrice","Hermione_mais","Rogue_dark","Dumbledore_intervention","Actes_illegaux"]].plot(
    kind="bar", stacked=False, figsize=(14,7), colormap='Set3'
)
plt.title("Événements spéciaux et actes répréhensibles par livre")
plt.ylabel("Occurrences")
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig("evenements_par_livre.png")
plt.show()

# 3️⃣ Normalisation par 100 pages (récurrence)
cols_100 = [col for col in stats_plot.columns if "_par100pages" in col]
plt.figure(figsize=(14,7))
stats_plot[cols_100].plot(kind="bar", stacked=False, figsize=(14,7), colormap='tab20')
plt.title("Occurrences normalisées par 100 pages")
plt.ylabel("Occurrences / 100 pages")
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig("stats_par100pages.png")
plt.show()

# 4️⃣ Graphiques de tendance par livre (ligne)
plt.figure(figsize=(14,7))
stats_plot[["Harry_cicatrice","Hermione_mais","Rogue_dark","Dumbledore_intervention","Actes_illegaux"]].plot(
    kind="line", marker='o', figsize=(14,7)
)
plt.title("Tendances des événements dans les livres")
plt.ylabel("Occurrences")
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig("tendances_evenements.png")
plt.show()


# -------------------------------------------------
# RAPPORT WORD AUTOMATIQUE
# -------------------------------------------------
doc = Document()
doc.add_heading("Rétrospective Harry Potter", 0)
doc.add_paragraph("Analyse objective des 7 livres de la saga Harry Potter. Les statistiques incluent les occurrences de différents événements et dialogues, normalisés par 100 pages pour comparaison.")

# Ajouter un tableau des stats par livre
doc.add_heading("Statistiques par livre", level=1)
table = doc.add_table(rows=1, cols=len(stats.columns))
hdr_cells = table.rows[0].cells
for i, col in enumerate(stats.columns):
    hdr_cells[i].text = col

for index, row in stats.iterrows():
    row_cells = table.add_row().cells
    for i, col in enumerate(stats.columns):
        row_cells[i].text = str(row[col])

# Ajouter images
doc.add_page_break()
doc.add_paragraph("Graphiques générés :")
doc.add_picture("stats_par_livre.png", width=Inches(6))
doc.add_picture("stats_par100pages.png", width=Inches(6))

doc.save("HarryPotter_Rapport.docx")
print("✅ Rapport Word généré : HarryPotter_Rapport.docx")
