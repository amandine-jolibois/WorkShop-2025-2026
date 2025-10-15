#!/usr/bin/env python3
# analyse_hp.py
import os, re, json, argparse
from collections import Counter, defaultdict
import pandas as pd
import matplotlib.pyplot as plt
from wordcloud import WordCloud
from tqdm import tqdm
import regex
import nltk
from nltk.tokenize import sent_tokenize
from docx import Document

nltk.download('punkt', quiet=True)

def load_text(path):
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()

def normalize_text(s):
    s = s.replace('\r\n','\n')
    return s

def count_phrases(text, phrases):
    lower = text.lower()
    return {p: lower.count(p.lower()) for p in phrases}

def approx_speech_counts(text, characters):
    quotes = regex.split(r'["“”«»\']', text)
    counts = Counter()
    for i, q in enumerate(quotes):
        if i % 2 == 1:
            snippet = q[:200]
            context = ''
            if i-1 >= 0: context += quotes[i-1][-200:]
            if i+1 < len(quotes): context += quotes[i+1][:200]
            assigned = False
            for c in characters:
                if re.search(r'\\b' + re.escape(c.lower()) + r'\\b', (context + snippet).lower()):
                    counts[c] += 1
                    assigned = True
                    break
            if not assigned:
                counts['UNKNOWN'] += 1
    return counts

def word_counts(text, top_n=50):
    words = regex.findall(r"\\p{L}+", text.lower())
    c = Counter(words)
    return c.most_common(top_n)

def plot_bar(d, title, outpath, xlabel='Count', ylabel='Item', horiz=False):
    items = list(d.keys())
    vals = list(d.values())
    plt.figure(figsize=(9, max(3, len(items)*0.3)))
    if horiz:
        plt.barh(items[::-1], vals[::-1])
        plt.xlabel(xlabel)
        plt.title(title)
    else:
        plt.bar(items, vals)
        plt.xticks(rotation=45, ha='right')
        plt.ylabel(ylabel)
        plt.title(title)
    plt.tight_layout()
    plt.savefig(outpath)
    plt.close()

def save_wordcloud(text, outpath):
    try:
        wc = WordCloud(width=1200, height=600, background_color='white').generate(text)
        wc.to_file(outpath)
    except Exception as e:
        # skip wordcloud if font issues
        with open(outpath + ".txt", "w", encoding="utf-8") as f:
            f.write("wordcloud generation skipped: " + str(e))

def main(books_dir, out_dir, page_counts):
    os.makedirs(out_dir, exist_ok=True)
    files = sorted([f for f in os.listdir(books_dir) if f.lower().endswith('.txt')])
    phrases = ["cicatrice", "mais", "dumbledore", "harry", "hermione", "ron", "severus", "rogue", "avada kedavra", "sang", "mort"]
    characters = ["Harry", "Hermione", "Ron", "Dumbledore", "Severus", "Snape", "Rogue", "Voldemort", "Draco", "Hagrid"]
    rows = []
    for fname in tqdm(files, desc="Analysing books"):
        key = os.path.splitext(fname)[0]
        text = load_text(os.path.join(books_dir, fname))
        text = normalize_text(text)
        n_words = len(regex.findall(r"\\p{L}+", text))
        phrase_counts = count_phrases(text, phrases)
        wc_top = word_counts(text, top_n=40)
        speech = approx_speech_counts(text, characters)
        pages = page_counts.get(key, None) if page_counts else None
        rows.append({
            "book": key,
            "words": n_words,
            "pages": pages or "",
            **{f"phrase_{p}": phrase_counts.get(p,0) for p in phrases},
            "speech_unknown": speech.get('UNKNOWN',0),
            **{f"speech_{c}": speech.get(c,0) for c in characters}
        })
        plot_bar(dict(wc_top), title=f"Top words - {key}", outpath=os.path.join(out_dir, f"{key}_topwords.png"), horiz=True)
        save_wordcloud(" ".join([w for w,_ in wc_top]), os.path.join(out_dir, f"{key}_wordcloud.png"))
    df = pd.DataFrame(rows)
    df.to_csv(os.path.join(out_dir, "summary.csv"), index=False)
    phrase_cols = [c for c in df.columns if c.startswith("phrase_")]
    ph = df.set_index('book')[phrase_cols].T
    ph.plot(kind='bar', figsize=(12,6))
    plt.title("Comparaison des phrases par livre")
    plt.tight_layout()
    plt.savefig(os.path.join(out_dir, "phrases_comparison.png"))
    plt.close()
    speech_cols = [c for c in df.columns if c.startswith("speech_")]
    s = df.set_index('book')[speech_cols].T
    s.plot(kind='bar', figsize=(12,6))
    plt.title("Estim. prises de parole par personnage (heuristique)")
    plt.tight_layout()
    plt.savefig(os.path.join(out_dir, "speech_comparison.png"))
    plt.close()
    doc = Document()
    doc.add_heading("Rapport automatique — Le Procès de J.K. Rowling", level=1)
    doc.add_paragraph("Ce document est généré automatiquement. Méthodologie : occurrences par recherche littérale, attribution de paroles par heuristique de guillemets/attribution.")
    doc.add_paragraph(f"Livres analysés : {len(files)}")
    doc.add_paragraph("Fichiers de sortie : summary.csv, graphiques PNG dans le dossier output.")
    doc.save(os.path.join(out_dir, "rapport_synthese.docx"))
    print("Done. Outputs saved in", out_dir)

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--books_dir", required=True)
    parser.add_argument("--out_dir", required=True)
    parser.add_argument("--page_counts", default=None)
    args = parser.parse_args()
    page_counts = {}
    if args.page_counts:
        try:
            with open(args.page_counts, 'r', encoding='utf-8') as f:
                page_counts = json.load(f)
        except:
            page_counts = {}
    main(args.books_dir, args.out_dir, page_counts)
